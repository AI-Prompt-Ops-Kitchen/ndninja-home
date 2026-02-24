from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from schemas.assembly import AssembledSOP


def generate_docx(assembled_sop: AssembledSOP) -> BytesIO:
    """Generate a DOCX file from assembled SOP data."""
    doc = Document()

    # Title
    title = doc.add_heading(assembled_sop.sop_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    table = doc.add_table(rows=3, cols=2)
    table.style = "Light Grid Accent 1"
    cells = table.rows[0].cells
    cells[0].text = "SOP Code"
    cells[1].text = assembled_sop.sop_code
    cells = table.rows[1].cells
    cells[0].text = "Standards"
    cells[1].text = assembled_sop.combo_key
    cells = table.rows[2].cells
    cells[0].text = "Content Tier"
    cells[1].text = assembled_sop.content_tier.title()

    doc.add_paragraph()

    # Sections
    for section in assembled_sop.sections:
        doc.add_heading(f"{section.section_number} {section.section_title}", level=1)
        for paragraph_text in section.body.split("\n\n"):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())

    # Traceability matrix (enhanced tier)
    if assembled_sop.traceability:
        doc.add_page_break()
        doc.add_heading("Traceability Matrix", level=1)
        t_table = doc.add_table(rows=1, cols=4)
        t_table.style = "Light Grid Accent 1"
        headers = t_table.rows[0].cells
        headers[0].text = "Standard"
        headers[1].text = "Clause"
        headers[2].text = "Requirement"
        headers[3].text = "Section"

        for entry in assembled_sop.traceability:
            row = t_table.add_row().cells
            row[0].text = entry.standard_code
            row[1].text = entry.requirement_clause
            row[2].text = entry.requirement_title
            row[3].text = entry.section_number

    # Cross-references
    if assembled_sop.cross_references:
        doc.add_heading("Cross-References", level=1)
        for ref in assembled_sop.cross_references:
            doc.add_paragraph(
                f"{ref['target_sop_code']} — {ref['target_sop_title']} ({ref['relationship_type']})",
                style="List Bullet",
            )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
